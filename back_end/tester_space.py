#python-docx
from docx import *

from os import walk
document = Document()
folder = "C:\\Users\\avino\\OneDrive\\Desktop\\8bit_scouting\\back_end\\servers\\"
# "C:\Users\avino\OneDrive\Desktop\8bit_scouting\back_end\servers"
filenames = next(walk(folder), (None, None, []))[2]  # [] if no file
print(filenames)
for file in filenames[::1]:
    p = document.add_paragraph()
    p.add_run(f'the file: {file}\n').bold = True
    file = open(f'{folder}{file}')
    #print(file.read())
    try:
        text = file.read()
        p.add_run(text)
        p.add_run('\n\n')
        file.close()
    except Exception as error:
        print(file)

document.save('demo.docx')
